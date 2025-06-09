from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, send_file
from shareplum import Site
from shareplum import Office365
from docx import Document
from io import BytesIO
import os
import json

# Inicializa a aplicação Flask
app = Flask(__name__)
app.secret_key = "your_secret_key"  # Chave secreta para mensagens flash

# Variáveis globais para armazenar as credenciais do SharePoint
username = None
password = None

# Rota de login, responsável por autenticar o usuário no SharePoint
@app.route("/", methods=["GET", "POST"])
def login():
    global username, password
    if request.method == "POST":
        username = request.form.get("email")
        password = request.form.get("password")

        # Verifica se os campos foram preenchidos
        if not username or not password:
            flash("Email e senha não podem estar vazios!", "error")
            return redirect(url_for("login"))

        try:
            # Tenta autenticar no SharePoint
            Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
            flash("Login bem-sucedido!", "success")
            return redirect(url_for("main"))
        except Exception as e:
            flash(f"Erro no login: {e}", "error")
            return redirect(url_for("login"))

    # Renderiza a página de login
    return render_template("login.html")

# Rota principal, exibe a lista de itens do SharePoint com filtros opcionais
@app.route("/main", methods=["GET", "POST"])
def main():
    try:
        # Autentica e conecta ao SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        # Inicializa filtros de busca
        status_filter = request.form.get("status_filter") if request.method == "POST" else None
        id_filter = request.form.get("id_filter") if request.method == "POST" else None

        # Monta a query de busca conforme os filtros
        query = {}
        if status_filter:
            query.setdefault('Where', []).append(('Eq', 'Status', status_filter))
        if id_filter:
            query.setdefault('Where', []).append(('Eq', 'ID', id_filter))

        # Busca os itens no SharePoint
        items = sp_list.GetListItems(query=query) if query else sp_list.GetListItems()
    except Exception as e:
        flash(f"Erro ao acessar a lista do SharePoint: {e}", "error")
        items = []

    # Lista de status para o filtro
    status_list = ["Edição", "Aprovado", "Invalidado"] 
    # Renderiza a página principal com os itens e filtros
    return render_template("main.html", items=items, status_list=status_list, status_filter=status_filter, id_filter=id_filter)

# Rota para criar um novo item no SharePoint (formulário)
@app.route("/form", methods=["GET", "POST"])
def form():
    linhas_json = []
    status_list = ["Edição", "Aprovado", "Invalidado"]  # Lista de status para o select
    if request.method == "POST":
        # Coleta os dados do formulário
        form_data = {
            'Status': request.form.get("status"),
            'Numero SEI': request.form.get("numero_sei"),
            'Nome': request.form.get("nome"),
            'Endereço': request.form.get("endereco"),
            'CPF/CNPJ': request.form.get("cpf_cnpj"),
            'Endereço Numero': request.form.get("endereco_numero"),
            'Bairro': request.form.get("bairro"),
            'UF': request.form.get("uf"),
            'CEP': request.form.get("cep"),
            'Telefone': request.form.get("telefone"),
        }

        # Processa os campos dinâmicos do formulário (intervenções ambientais)
        tipo_intervencao_list = request.form.getlist('tipo_intervencao[]')
        quantidade_list = request.form.getlist('quantidade[]')
        unidade_list = request.form.getlist('unidade[]')
        linhas = []
        for tipo, qtd, un in zip(tipo_intervencao_list, quantidade_list, unidade_list):
            linhas.append({
                'tipo_intervencao': tipo,
                'quantidade': qtd,
                'unidade': un
            })
        form_data['JSON'] = json.dumps(linhas, ensure_ascii=False)

        # Validação dos campos obrigatórios
        if not form_data['Status'] or not form_data['Numero SEI']:
            flash("Os campos 'Status' e 'Número SEI' são obrigatórios!", "error")
            return redirect(url_for("form"))

        try:
            # Autentica e conecta ao SharePoint
            authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
            site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
            sp_list = site.List('Base de Dados')

            # Insere o novo item no SharePoint
            sp_list.UpdateListItems(data=[form_data], kind='New')
            flash("Item inserido com sucesso no SharePoint!", "success")
        except Exception as e:
            flash(f"Erro ao inserir item: {e}", "error")

        return redirect(url_for("form"))

    # Se for GET, pode carregar um item existente para edição (não implementado aqui)
    item = None # ou busque o item do SharePoint se necessário
    if item and item.get('JSON'):
        try:
            linhas_json = json.loads(item['JSON'])
        except Exception:
            linhas_json = []
    # Renderiza o formulário
    return render_template("form.html", item=item, status_list=status_list, linhas_json=linhas_json)

# Rota para editar um item existente do SharePoint
@app.route("/edit/<item_id>", methods=["GET", "POST"])
def edit(item_id):
    try:
        # Autentica e conecta ao SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        print(f"Editando item com ID: {item_id}")

        if request.method == "POST":
            # Coleta os dados atualizados do formulário
            form_data = {
                'ID': item_id,
                'Status': request.form.get("status"),
                'Numero SEI': request.form.get("numero_sei"),
                'Nome': request.form.get("nome"),
                'Endereço': request.form.get("endereco"),
                'CPF/CNPJ': request.form.get("cpf_cnpj"),
                'Endereço Numero': request.form.get("endereco_numero"),
                'Bairro': request.form.get("bairro"),
                'UF': request.form.get("uf"),
                'CEP': request.form.get("cep"),
                'Telefone': request.form.get("telefone"),
            }
            # Processa os campos dinâmicos do formulário
            tipo_intervencao_list = request.form.getlist('tipo_intervencao[]')
            quantidade_list = request.form.getlist('quantidade[]')
            unidade_list = request.form.getlist('unidade[]')
            linhas = []
            for tipo, qtd, un in zip(tipo_intervencao_list, quantidade_list, unidade_list):
                linhas.append({
                    'tipo_intervencao': tipo,
                    'quantidade': qtd,
                    'unidade': un
                })
            form_data['JSON'] = json.dumps(linhas, ensure_ascii=False)

            # Atualiza o item existente no SharePoint
            sp_list.UpdateListItems(data=[form_data], kind='Update')
            flash("Item atualizado com sucesso no SharePoint!", "success")
            return redirect(url_for("main"))

        # Busca o item para preencher o formulário de edição
        item = sp_list.GetListItems(
            fields=['ID', 'Status', 'Numero SEI', 'Nome', 'Endereço', 'CPF/CNPJ', 'Endereço Numero', 'Bairro', 'UF', 'CEP', 'Telefone', 'JSON'],
            query={'Where': [('Eq', 'ID', item_id)]}
        )
        if not item:
            flash("Item não encontrado!", "error")
            return redirect(url_for("main"))

        item = item[0]
        # Carrega o JSON para os campos dinâmicos
        linhas_json = []
        if item.get('JSON'):
            try:
                linhas_json = json.loads(item['JSON'])
            except Exception:
                linhas_json = []

        status_list = ["Edição", "Aprovado", "Invalidado"]
        # Renderiza o formulário de edição
        return render_template("form.html", status_list=status_list, item=item, linhas_json=linhas_json)
    except Exception as e:
        flash(f"Erro ao acessar ou atualizar o item: {e}", "error")
        return redirect(url_for("main"))

# Rota para gerar e baixar um documento Word preenchido com os dados do item selecionado
@app.route("/download/<item_id>", methods=["GET"])
def download(item_id):
    try:
        # Autentica e conecta ao SharePoint
        authcookie = Office365('https://meioambientemg.sharepoint.com', username=username, password=password).GetCookies()
        site = Site('https://meioambientemg.sharepoint.com/sites/BasedeDados', authcookie=authcookie)
        sp_list = site.List('Base de Dados')

        print(f"Gerando download para o item com ID: {item_id}")

        # Busca os dados do item
        item = sp_list.GetListItems(
            fields=['Nome', 'Endereço', 'Telefone', 'JSON'],
            query={'Where': [('Eq', 'ID', item_id)]}
        )
        if not item:
            print("DEBUG: Item não encontrado!")
            flash("Item não encontrado!", "error")
            return redirect(url_for("main"))

        item = item[0]

        # Carrega o template Word
        template_path = os.path.join(os.path.dirname(__file__), "2Modelo Parecer (Fabio) 2.docx")
        if not os.path.exists(template_path):
            print(f"DEBUG: Template não encontrado em {template_path}")
            flash("Template Word não encontrado!", "error")
            return redirect(url_for("main"))

        print(f"DEBUG: Template encontrado em {template_path}")
        doc = Document(template_path)

        # Adiciona os dados do JSON na tabela existente do documento Word
        if item.get("JSON"):
            try:
                dados_tabela = json.loads(item["JSON"])
                # Supondo que o JSON é uma lista de dicts com as chaves: tipo_intervencao, quantidade, unidade
                table = doc.tables[0]  # O documento é inteiramente uma unica tabela
                row_index = 18  # Índice da linha onde começa a inserir (ajuste conforme seu modelo)
                for i, linha in enumerate(dados_tabela):
                    new_row = table.add_row()
                    # Move a nova linha para logo após a linha desejada
                    table._tbl.remove(new_row._tr)
                    table._tbl.insert(row_index + 2 + i, new_row._tr)
                    # Preenche as células (ajuste os índices conforme sua tabela)
                    new_row.cells[0].text = linha.get("tipo_intervencao", "")
                    new_row.cells[3].text = linha.get("quantidade", "")
                    new_row.cells[8].text = linha.get("unidade", "")
                    # Mescla células porque quando a linha é adicionada ela vem com 12 colunas.
                    new_row.cells[0].merge(new_row.cells[2])
                    new_row.cells[3].merge(new_row.cells[7])
                    new_row.cells[8].merge(new_row.cells[12])
            except Exception as e:
                doc.add_paragraph("Erro ao processar dados de intervenção ambiental.")

        # Função para substituir texto em parágrafos e células de tabelas
        def replace_text_in_paragraphs(paragraphs, replacements):
            for paragraph in paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)

        # Função para substituir texto em tabelas
        def replace_text_in_tables(tables, replacements):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_paragraphs(cell.paragraphs, replacements)

        # Dicionário de substituições para os campos do documento
        replacements = {
            "{{Nome}}": item.get("Nome", ""),
            "{{Endereço}}": item.get("Endereço", ""),
            "{{Telefone}}": item.get("Telefone", "")
        }

        # Substitui os campos nos parágrafos do corpo do documento
        replace_text_in_paragraphs(doc.paragraphs, replacements)

        # Substitui os campos nas tabelas
        replace_text_in_tables(doc.tables, replacements)

        # Substitui os campos nos cabeçalhos e rodapés
        for section in doc.sections:
            replace_text_in_paragraphs(section.header.paragraphs, replacements)
            replace_text_in_paragraphs(section.footer.paragraphs, replacements)

        # Salva o documento em memória e envia para download
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Parecer_{item_id}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("DEBUG: Exception capturada no download:", e)
        flash(f"Erro ao gerar o arquivo Word: {e}", "error")
        return redirect(url_for("main"))

# Inicia a aplicação Flask (lembrar de retirar o debug=True em produção)
if __name__ == "__main__":
    app.run(debug=True)
